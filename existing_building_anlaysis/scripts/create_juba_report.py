from pathlib import Path
from datetime import date

import json
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
REPORT = OUT / "Juba_Building_Data_Consistency_Assessment.docx"

NAVY = "1F4D78"
BLUE = "2E74B5"
TEAL = "1D7A75"
PALE_BLUE = "EAF2F8"
PALE_TEAL = "E8F4F2"
PALE_GOLD = "FFF4D6"
LIGHT = "F2F4F7"
MID = "D0D7DE"
DARK = "243447"
MUTED = "667085"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:" + key), str(value))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_col_widths(table, widths):
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)


def style_table(table, header_fill=LIGHT, font_size=9, widths=None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    if widths:
        set_col_widths(table, widths)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(
                cell,
                bottom={"val": "single", "sz": "4", "color": MID},
            )
            if r_idx == 0:
                set_cell_shading(cell, header_fill)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = RGBColor.from_string(DARK)
                    if r_idx == 0:
                        run.bold = True
        if r_idx == 0:
            set_repeat_table_header(row)


def add_table(doc, headers, rows, widths=None, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = str(h)
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = str(value)
    style_table(table, widths=widths, font_size=font_size)
    return table


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, sep, text, end])


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Juba building data assessment  |  ")
    r.font.name = "Calibri"
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    add_field(p, "PAGE")


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 10),
        ("Subtitle", 15, MUTED, 0, 12),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(7)
    caption.paragraph_format.keep_with_next = False

    header = section.header
    p = header.paragraphs[0]
    p.text = "JUBA BUILDING DATA CONSISTENCY ASSESSMENT"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.runs[0]
    r.font.name = "Calibri"
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_number(section)

    core = doc.core_properties
    core.title = "Juba Building Data Consistency Assessment"
    core.subject = "Raster coverage, building heights, neighborhood summaries, and footprint geometry comparisons"
    core.author = "Building data comparison analysis"
    core.keywords = "Juba, building footprints, building heights, Overture, Google 2.5D, Global Building Atlas, 3D-GloBFP, TEMPO, WSF"
    return doc


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)
    return p


def add_figure(doc, filename, caption, width=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(OUT / filename), width=Inches(width))
    add_caption(doc, caption)


def add_callout(doc, title, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, start=180, bottom=130, end=180)
    set_cell_border(cell, left={"val": "single", "sz": "16", "color": TEAL})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.05
    for run in p2.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.add_run(item)


def page_break(doc):
    doc.add_page_break()


def build_report():
    footprint = pd.read_csv(OUT / "juba_30m_source_summary.csv")
    footprint_sensitivity = pd.read_csv(OUT / "juba_30m_threshold_sensitivity.csv")
    height = pd.read_csv(OUT / "juba_height_source_summary.csv")
    height100 = height[height.grid == "100m"].set_index("source")
    height_pairs = pd.read_csv(OUT / "juba_height_pairwise_agreement.csv")
    height_sensitivity = pd.read_csv(OUT / "juba_height_sensitivity.csv").set_index("scenario")
    neighborhood = pd.read_csv(OUT / "juba_neighborhood_overview.csv")
    neighborhood_meta = json.loads((OUT / "juba_neighborhood_analysis_metadata.json").read_text())
    geometry_meta = json.loads((OUT / "juba_geometry_analysis_metadata.json").read_text())
    geometry = pd.read_csv(OUT / "juba_geometry_overall_summary.csv")

    def height_pair(grid, left, right):
        return height_pairs[
            (height_pairs.grid == grid)
            & (height_pairs.source_a == left)
            & (height_pairs.source_b == right)
        ].iloc[0]

    doc = setup_document()

    # Cover — editorial_cover pattern.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("TECHNICAL ASSESSMENT")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(TEAL)
    r.font.letter_spacing = Pt(1.2)

    doc.add_paragraph("Juba Building Data\nConsistency Assessment", style="Title")
    doc.add_paragraph(
        "Raster coverage, building heights, neighborhood summaries, and footprint geometry comparisons",
        style="Subtitle",
    )
    line = doc.add_paragraph()
    line.paragraph_format.space_after = Pt(14)
    run = line.add_run("Juba, South Sudan  |  August 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    add_figure(
        doc,
        "juba_30m_rgb_composite.png",
        "Cover image. RGB composite of the first three 30 m comparison bands; colors indicate where sources agree or differ in mapped building fraction.",
        width=6.25,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Study extent")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run(
        f"  Analysis AOI: {neighborhood_meta['analyzed_area_km2']:.2f} km²  •  "
        f"Local reporting units: {neighborhood_meta['juba_features']:,}"
    )

    page_break(doc)

    doc.add_heading("Executive summary", level=1)
    doc.add_paragraph(
        "This assessment compares six global building datasets over Juba using a common raster framework and adds GlobalBuildingAtlas GBA.Height as a fifth modeled height product. It summarizes consistency by local reporting unit and directly compares vector footprint geometry where polygons are available. The results describe inter-product consistency and relative completeness, not accuracy against independent ground truth."
    )
    add_callout(
        doc,
        "Main conclusion",
        "The three vector-footprint products are broadly complete relative to the multi-source consensus, while Google 2.5D is less complete under the selected 30 m threshold. GBA.Height is systematically lower than Google 2.5D and 3D-GloBFP over Juba and increases measured inter-product spread. It is not independent of TEMPO because both use PlanetScope imagery, and its footprint lineage overlaps existing Google, Microsoft, and OSM-derived products.",
        fill=PALE_TEAL,
    )
    doc.add_heading("Key findings", level=2)
    add_bullets(
        doc,
        [
            "At 30 m, the completeness proxy is 97.8% for Global Building Atlas, 97.1% for Overture, 95.8% for 3D-GloBFP, and 71.4% for Google 2.5D using a 25 m² positive-area threshold.",
            "GBA.Height has 19,611 valid 100 m cells (25.0% of the AOI), a median of 1.75 m, and a 90th percentile of 2.91 m after restriction to GBA footprints.",
            "Including GBA.Height raises the median 100 m inter-product range from 2.10 m to 2.69 m; the corresponding 90th-percentile range rises from 3.74 m to 4.08 m.",
            "At 100 m, GBA.Height is closest to TEMPO by MAE (0.81 m), while it differs from 3D-GloBFP by 2.30 m MAE. Shared imagery means the TEMPO agreement is not independent validation.",
            f"Across {neighborhood_meta['juba_features']:,} local reporting units, the median height range is 2.33 m without GBA and 2.84 m with GBA. The largest included-product range is 16.66 m in unit CN-38.",
            "Global Building Atlas and 3D-GloBFP have a median neighborhood union IoU of 0.974 and a median one-to-one object IoU of approximately 1.0, indicating extensive shared or duplicated footprint geometry.",
        ],
    )
    doc.add_heading("Decision use", level=2)
    add_table(
        doc,
        ["Question", "Recommended use"],
        [
            ["Where are building-presence gaps?", "Use the 30 m consensus-gap maps, then inspect neighborhood summaries."],
            ["Which source is most complete?", "Use the completeness proxy, but account for correlated inputs and thresholds."],
            ["Which height surface is most reliable?", "Treat agreement as evidence of consistency, not accuracy; validate priority areas independently."],
            ["Which polygons align geometrically?", "Use neighborhood union IoU and object-level matching together."],
        ],
        widths=[2.05, 4.25],
        font_size=9,
    )

    page_break(doc)

    doc.add_heading("1. Data and analytical framework", level=1)
    doc.add_paragraph(
        "The analysis combines gridded and polygon building products by rasterizing or aggregating them to common grids. Two spatial units are used: regular cells for city-wide comparisons and neighborhood polygons (ID_SEG) for local diagnostics."
    )
    doc.add_heading("Sources and analytical roles", level=2)
    add_table(
        doc,
        ["Source", "Footprint / built-up role", "Height role"],
        [
            ["TEMPO", "100 m built fraction", "100 m height"],
            ["Overture", "Vector building footprints", "Very sparse attributes; excluded from height comparison"],
            ["Google 2.5D", "Modeled building fraction / presence", "Height at 30 m and 100 m"],
            ["Global Building Atlas", "Vector building footprints", "GBA.Height at 30 m and 100 m; footprint-masked"],
            ["3D-GloBFP", "Vector building footprints", "Height at 30 m and 100 m"],
            ["WSF 3D v2", "100 m building fraction", "100 m height"],
        ],
        widths=[1.55, 2.25, 2.5],
        font_size=8.7,
    )
    doc.add_heading("Core metrics", level=2)
    add_bullets(
        doc,
        [
            "Building coverage: mapped building area divided by analysis-unit area.",
            "Completeness proxy: percentage of multi-source consensus cells in which a source is positive. For the primary 30 m result, a cell is positive when mapped building area is at least 25 m².",
            "Height agreement: common-cell bias, mean absolute error (MAE), root mean squared error (RMSE), rank correlation, and height-class agreement.",
            "Neighborhood union IoU: intersection divided by union after merging all source footprints within a neighborhood.",
            "Object matching: polygons are linked when overlap is at least 1 m² and at least 10% of the smaller footprint. One-to-one matches are then assessed with IoU, centroid distance, Hausdorff distance, and orientation difference.",
        ],
    )
    add_callout(
        doc,
        "Interpretation guardrail",
        "A consensus completeness score measures agreement with the other sources, not absolute truth. GBA.Height is excluded from independent-source counts: it shares PlanetScope imagery with TEMPO, while the GBA LoD1/footprint chain incorporates Google, Microsoft, OSM, and other fused inputs. Temporal differences, scale, thresholds, and shared inputs can all produce systematic agreement or disagreement.",
        fill=PALE_GOLD,
    )

    page_break(doc)

    doc.add_heading("2. Raster footprint coverage at 30 m", level=1)
    doc.add_paragraph(
        "The 30 m comparison is the most spatially detailed city-wide footprint assessment. It uses 135,232 consensus-positive cells under the primary 25 m² threshold."
    )
    add_figure(
        doc,
        "juba_30m_source_fractions.png",
        "Figure 1. Estimated building fraction by source on the 30 m grid. Differences reflect both mapped extent and source-specific representation of building area.",
        width=6.15,
    )
    add_table(
        doc,
        ["Source", "Positive cells", "Completeness proxy", "Gap cells", "Built area"],
        [[
            row.source,
            f"{row.positive_30m_cells:,.0f}",
            f"{row.consensus_recall_proxy_pct:.1f}%",
            f"{row.gap_30m_cells:,.0f}",
            f"{row.estimated_built_area_km2:.2f} km²",
        ] for row in footprint.itertuples()],
        widths=[1.8, 1.05, 1.35, 0.9, 1.2],
        font_size=8.8,
    )
    add_caption(doc, "Table 1. Primary 30 m footprint summary using the 25 m² building-area threshold.")

    page_break(doc)

    doc.add_heading("2.1 Spatial gaps and threshold sensitivity", level=2)
    doc.add_paragraph(
        "The gap maps identify consensus-positive cells that are absent from each source. Google 2.5D shows the broadest and most spatially continuous gaps, while the vector products have smaller, more localized omissions."
    )
    add_figure(
        doc,
        "juba_30m_source_gaps.png",
        "Figure 2. Source-specific gaps against the multi-source 30 m consensus. A marked cell is positive in the consensus but absent from the named source.",
        width=6.05,
    )
    doc.add_heading("Sensitivity to the positive-cell threshold", level=3)
    add_table(
        doc,
        ["Source", "10 m²", "25 m²", "50 m²"],
        [[source] + [
            f"{footprint_sensitivity[(footprint_sensitivity.source == source) & (footprint_sensitivity.building_area_threshold_m2 == threshold)].iloc[0].consensus_recall_proxy_pct:.1f}%"
            for threshold in (10.0, 25.0, 50.0)
        ] for source in ("Overture", "Google 2.5D", "GlobalBuildingAtlas", "3D-GloBFP")],
        widths=[2.4, 1.3, 1.3, 1.3],
        font_size=9,
    )
    add_caption(doc, "Table 2. Completeness proxy under alternative building-area thresholds.")

    page_break(doc)

    doc.add_heading("3. Building height comparison", level=1)
    doc.add_paragraph(
        "Five modeled height products were evaluated at the common 100 m support; Google 2.5D, GBA.Height, and 3D-GloBFP were also compared at 30 m. GBA.Height is a nominal 3 m PlanetScope-derived raster. It was restricted to Global Building Atlas footprints and aggregated by valid building-pixel area. Overture height attributes remain too sparse for a defensible surface-level comparison."
    )
    add_figure(
        doc,
        "juba_100m_height_comparison.png",
        "Figure 3. Mean building height surfaces at 100 m for the five modeled products, including footprint-masked GBA.Height.",
        width=6.1,
    )
    add_table(
        doc,
        ["Source", "Valid AOI", "Median", "90th pct.", "Est. volume"],
        [[
            source,
            f"{height100.loc[source].height_coverage_of_aoi_pct:.1f}%",
            f"{height100.loc[source].median_height_m:.2f} m",
            f"{height100.loc[source].p90_height_m:.2f} m",
            f"{height100.loc[source].estimated_built_volume_m3 / 1e6:.1f} million m³",
        ] for source in ("TEMPO", "Google 2.5D", "GBA.Height", "3D-GloBFP", "WSF 3D v2")],
        widths=[1.55, 1.0, 1.05, 1.05, 1.65],
        font_size=8.8,
    )
    add_caption(doc, "Table 3. Height distribution and coverage summary on the 100 m grid.")

    page_break(doc)

    doc.add_heading("3.1 Height agreement and uncertainty", level=2)
    add_figure(
        doc,
        "juba_100m_height_diagnostics.png",
        "Figure 4. Valid-product counts and inter-product height ranges with GBA excluded and included. GBA is excluded from any independent-source interpretation.",
        width=5.45,
    )
    add_table(
        doc,
        ["Comparison", "Bias A−B", "MAE", "RMSE", "Rank corr.", "Same class"],
        [[
            f"{left} vs {right}, {grid}",
            f"{height_pair(grid, left, right).mean_bias_a_minus_b_m:.2f} m",
            f"{height_pair(grid, left, right).mae_m:.2f} m",
            f"{height_pair(grid, left, right).rmse_m:.2f} m",
            f"{height_pair(grid, left, right).spearman_height:.3f}",
            f"{height_pair(grid, left, right).same_height_class_pct:.1f}%",
        ] for grid, left, right in (
            ("100m", "TEMPO", "GBA.Height"),
            ("100m", "Google 2.5D", "GBA.Height"),
            ("100m", "GBA.Height", "3D-GloBFP"),
            ("100m", "Google 2.5D", "3D-GloBFP"),
            ("30m", "Google 2.5D", "GBA.Height"),
            ("30m", "GBA.Height", "3D-GloBFP"),
        )],
        widths=[2.15, 0.82, 0.72, 0.72, 0.85, 0.85],
        font_size=7.5,
    )
    add_caption(doc, "Table 4. Selected pairwise height agreement statistics.")
    add_callout(
        doc,
        "Interpretation",
        "GBA.Height is about 2 m lower than Google 2.5D and 3D-GloBFP on average. Its closer agreement with TEMPO is not independent validation because both use PlanetScope imagery. Including GBA raises the median cell range from 2.10 m to 2.69 m.",
        fill=PALE_BLUE,
    )

    doc.add_heading("4. Neighborhood-level footprint consistency", level=1)
    doc.add_paragraph(
        f"The local analysis uses {neighborhood_meta['juba_features']:,} non-overlapping reporting units covering {neighborhood_meta['analyzed_area_km2']:.2f} km². Summaries help distinguish widespread source behavior from localized data gaps; many peripheral units contain no mapped buildings, so medians below are reported over units with valid values."
    )
    add_figure(
        doc,
        "juba_neighborhood_footprint_completeness.png",
        "Figure 5. Footprint completeness proxy by neighborhood for Overture, Google 2.5D, Global Building Atlas, and 3D-GloBFP.",
        width=6.05,
    )
    add_table(
        doc,
        ["Source", "Median completeness", "Median building coverage"],
        [
            ["Global Building Atlas", "100.00%", "0.64%"],
            ["Overture", "99.78%", "1.46%"],
            ["3D-GloBFP", "99.96%", "0.65%"],
            ["Google 2.5D", "43.00%", "0.55%"],
        ],
        widths=[2.45, 1.85, 2.0],
        font_size=9,
    )
    add_caption(doc, "Table 5. Median completeness over consensus-bearing units and median building coverage over occupied units.")

    page_break(doc)

    doc.add_heading("4.1 Neighborhood heights", level=2)
    add_figure(
        doc,
        "juba_neighborhood_mean_heights.png",
        "Figure 6. Building-area-weighted mean height by local reporting unit for five modeled products.",
        width=6.05,
    )
    add_table(
        doc,
        ["Source", "Median neighborhood mean height"],
        [[source, f"{neighborhood[f'height_{slug}_m'].median():.2f} m"] for source, slug in (
            ("3D-GloBFP", "globfp3d"), ("Google 2.5D", "google25d"),
            ("WSF 3D v2", "wsf3d"), ("TEMPO", "tempo"), ("GBA.Height", "gba_height")
        )],
        widths=[3.3, 3.0],
        font_size=9,
    )
    add_caption(doc, "Table 6. Median of neighborhood-level mean heights.")
    doc.add_paragraph(
        "Among units with comparable heights, the median range is 2.33 m with GBA excluded and 2.84 m with GBA included. This systematic vertical disagreement is large enough to affect floor-area, exposure, and built-volume estimates."
    )

    page_break(doc)

    doc.add_heading("4.2 Local disagreement hotspots", level=2)
    add_figure(
        doc,
        "juba_neighborhood_disagreement.png",
        "Figure 7. Local footprint-coverage range plus mean-height ranges with GBA excluded and included.",
        width=6.25,
    )
    add_callout(
        doc,
        "Priority review segments",
        "Unit CV-99 has the largest footprint-coverage range (23.39 percentage points). Unit CN-38 has the largest mean-height range (16.66 m), driven by a Google 2.5D mean of 19.05 m versus 2.39 m for TEMPO; GBA.Height is 4.89 m there. These are priority candidates for independent imagery or field review.",
        fill=PALE_GOLD,
    )
    doc.add_heading("Neighborhood interpretation", level=3)
    add_bullets(
        doc,
        [
            "Google 2.5D is the least complete source in 4,336 of 4,864 consensus-bearing reporting units under the selected threshold.",
            "Coverage disagreement is not limited to the urban fringe; several compact inner neighborhoods also show large spreads.",
            "A high height range can reflect genuine tall-building clusters, model differences, or misalignment between footprint fraction and height support.",
            "GBA.Height has no African training or validation samples documented in the release paper, so domain shift is an important Juba-specific limitation."
        ],
    )

    page_break(doc)

    doc.add_heading("5. Vector footprint geometry comparison", level=1)
    doc.add_paragraph(
        f"Three polygon sources were directly compared within the {geometry_meta['neighborhoods']:,} reporting units: {geometry_meta['feature_counts']['Overture']:,} Overture footprints, {geometry_meta['feature_counts']['GlobalBuildingAtlas']:,} Global Building Atlas footprints, and {geometry_meta['feature_counts']['3D-GloBFP']:,} 3D-GloBFP footprints. Union metrics describe total mapped shape; object metrics describe individual-building alignment and split/merge behavior."
    )
    add_figure(
        doc,
        "juba_geometry_neighborhood_union_iou.png",
        "Figure 8. Neighborhood union IoU for each vector-source pair. Higher values indicate greater overlap in the combined footprint area.",
        width=6.25,
    )
    add_figure(
        doc,
        "juba_geometry_neighborhood_object_iou.png",
        "Figure 9. Median one-to-one object IoU by neighborhood. This isolates matched individual footprints from total mapped extent.",
        width=6.25,
    )

    page_break(doc)

    doc.add_heading("5.1 Match structure and pairwise results", level=2)
    add_figure(
        doc,
        "juba_geometry_match_type_shares.png",
        "Figure 10. Shares of one-to-one, split/merge, many-to-many, and unmatched geometry groups by source pair.",
        width=6.25,
    )
    add_table(
        doc,
        ["Pair", "Union IoU", "Object IoU", "Centroid", "Hausdorff", "Matched A / B"],
        [
            ["Overture–GBA", "0.350", "0.585", "1.09 m", "1.64 m", "67.0% / 84.6%"],
            ["Overture–3D-GloBFP", "0.362", "0.592", "1.08 m", "1.62 m", "64.5% / 91.7%"],
            ["GBA–3D-GloBFP", "0.974", "≈1.000", "≈0.00 m", "≈0.00 m", "100% / 100%"],
        ],
        widths=[1.6, 0.85, 0.85, 0.85, 0.85, 1.3],
        font_size=8.2,
    )
    add_caption(doc, "Table 7. Median neighborhood and one-to-one geometry metrics; matched percentages are directional.")
    add_callout(
        doc,
        "Critical dependency finding",
        "Global Building Atlas and 3D-GloBFP are almost identical for matched one-to-one footprints: median object IoU is approximately 1.0, with essentially zero centroid and boundary displacement. Their agreement is therefore strong evidence of shared geometry, not two independent observations of the built environment.",
        fill=PALE_GOLD,
    )

    page_break(doc)

    doc.add_heading("6. Conclusions and recommended use", level=1)
    doc.add_heading("What the evidence supports", level=2)
    add_bullets(
        doc,
        [
            "For footprint presence, Overture, Global Building Atlas, and 3D-GloBFP provide broadly complete coverage relative to the multi-source consensus over Juba.",
            "For mapped building area, Overture is highest city-wide, while Global Building Atlas and 3D-GloBFP are closer to each other; the choice can materially change exposure denominators.",
            "For height, Google 2.5D and 3D-GloBFP remain the most consistent pair. GBA.Height is substantially lower than both, closer to TEMPO, and increases overall spread; none should be assumed accurate without independent validation.",
            "Neighborhood units are effective for prioritizing manual review because they preserve local context and make systematic source behavior visible.",
        ],
    )
    doc.add_heading("Recommended workflow for scaling up", level=2)
    add_table(
        doc,
        ["Stage", "Recommended action"],
        [
            ["1. Screen", "Run the 30 m footprint and 100 m height comparisons for each city using fixed thresholds and common metadata."],
            ["2. Localize", "Aggregate results to stable neighborhood or administrative polygons and rank footprint- and height-disagreement hotspots."],
            ["3. De-duplicate evidence", "Flag source pairs with near-identical geometry so correlated products do not receive independent weight."],
            ["4. Validate", "Review priority segments against recent imagery or authoritative local data before making accuracy claims."],
            ["5. Publish", "Retain raster layers, neighborhood tables, geometry match groups, and a concise city report for reproducibility."],
        ],
        widths=[1.25, 5.05],
        font_size=9,
    )
    doc.add_heading("Limitations", level=2)
    doc.add_paragraph(
        "The products represent different dates, definitions, minimum mapping units, and processing chains. The consensus is relative and may inherit shared omissions. GBA.Height is predominantly based on 2019 PlanetScope imagery with 2018 supplementation, is licensed CC BY-NC 4.0, and has no African training or validation samples documented in the release paper. It shares PlanetScope signal with TEMPO, while its footprint/LoD1 lineage overlaps Google, Microsoft, and OSM-derived inputs. Height products may encode non-equivalent vertical quantities. Rasterization, clipping, and reporting-unit boundaries introduce additional scale effects."
    )
    add_callout(
        doc,
        "Bottom line",
        "Use the combined analysis to find gaps, prioritize review, and understand sensitivity. Use an independent reference—recent imagery, local cadastral data, or field observations—when selecting a definitive footprint or height product for operational decisions.",
        fill=PALE_TEAL,
    )

    doc.add_heading("Appendix A. Reproducible output inventory", level=1)
    doc.add_paragraph(
        "The report summarizes the following principal machine-readable outputs. All are stored in the project outputs directory."
    )
    add_table(
        doc,
        ["Topic", "Primary files"],
        [
            ["30 m raster comparison", "juba_30m_comparison.tif; juba_30m_source_summary.csv; juba_30m_threshold_sensitivity.csv"],
            ["100 m height comparison", "juba_100m_height_comparison.tif; height source, pairwise, sensitivity, and hotspot tables"],
            ["GBA.Height provenance", "data/raw/gba_height/juba_download_manifest.json; four selectively extracted 0.2° GeoTIFFs"],
            ["Neighborhood analysis", "juba_neighborhood_consistency.gpkg; juba_neighborhood_overview.csv; footprint and height summary tables"],
            ["Geometry comparison", "juba_geometry_neighborhood_summary.gpkg; juba_geometry_overall_summary.csv; neighborhood union and object matching tables"],
            ["Review geometries", "Pair-specific review GeoPackages and full geometry match-group Parquet files"],
        ],
        widths=[1.75, 4.55],
        font_size=8.6,
    )
    doc.add_heading("Metric definitions", level=2)
    add_table(
        doc,
        ["Metric", "Meaning"],
        [
            ["Completeness proxy", "Recall against a multi-source consensus-positive mask."],
            ["Coverage range", "Maximum minus minimum building-coverage percentage across sources in an analysis unit."],
            ["Height range", "Maximum minus minimum source mean height within an analysis unit."],
            ["IoU", "Intersection area divided by union area; 1 indicates identical geometry."],
            ["Dice", "Twice the intersection divided by the sum of the two areas."],
            ["Hausdorff distance", "Maximum boundary separation between a matched pair; lower is better."],
        ],
        widths=[1.75, 4.55],
        font_size=8.8,
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("End of report")
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string(MUTED)

    REPORT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    build_report()
